# Create the document processor
document_processor_content = '''"""
Multi-format document processing with OCR support
"""
import os
import logging
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
from dataclasses import dataclass
import hashlib
import mimetypes

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
    from docx import Document as DocxDocument
    from bs4 import BeautifulSoup
    import pytesseract
    from PIL import Image
    import pdf2image
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")

from ..utils.text_processing import text_processor
from ..utils.chunking import chunking_manager
from ..config.settings import settings

@dataclass
class DocumentMetadata:
    """Document metadata container"""
    filename: str
    file_path: str
    file_size: int
    file_type: str
    mime_type: str
    created_at: str
    modified_at: str
    hash_md5: str
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None

@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    content: str
    metadata: DocumentMetadata
    chunks: List[Any]  # Will be Chunk objects
    processing_info: Dict[str, Any]

class DocumentProcessor:
    """Multi-format document processor with OCR support"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.md': self._process_markdown,
        }
        self.setup_logging()
    
    def setup_logging(self):
        """Setup logging for document processing"""
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(getattr(logging, settings.log_level))
        
        if not self.logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            self.logger.addHandler(handler)
    
    def process_document(self, file_path: Union[str, Path], 
                        chunking_strategy: str = 'recursive',
                        **chunking_kwargs) -> ProcessedDocument:
        """
        Process a document and extract text content
        
        Args:
            file_path: Path to the document file
            chunking_strategy: Strategy for chunking the document
            **chunking_kwargs: Additional arguments for chunking
            
        Returns:
            ProcessedDocument object
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file metadata
        metadata = self._extract_metadata(file_path)
        
        # Determine file type and process
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        self.logger.info(f"Processing document: {file_path}")
        
        try:
            # Process the document
            processing_func = self.supported_formats[file_extension]
            raw_content = processing_func(file_path)
            
            # Clean the text
            cleaned_content = text_processor.clean_text(raw_content)
            
            # Remove headers and footers
            cleaned_content = text_processor.remove_headers_footers(cleaned_content)
            
            # Update metadata with processing info
            metadata.word_count = len(cleaned_content.split())
            
            # Chunk the document
            document_id = metadata.hash_md5
            chunks = chunking_manager.chunk_document(
                cleaned_content, 
                document_id,
                chunking_strategy,
                **chunking_kwargs
            )
            
            processing_info = {
                'original_length': len(raw_content),
                'cleaned_length': len(cleaned_content),
                'chunk_count': len(chunks),
                'chunking_strategy': chunking_strategy,
                'processing_method': file_extension[1:]  # Remove the dot
            }
            
            self.logger.info(f"Successfully processed {file_path}: "
                           f"{len(chunks)} chunks, {metadata.word_count} words")
            
            return ProcessedDocument(
                content=cleaned_content,
                metadata=metadata,
                chunks=chunks,
                processing_info=processing_info
            )
            
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            raise
    
    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract file metadata"""
        stat = file_path.stat()
        
        # Calculate MD5 hash
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        return DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size=stat.st_size,
            file_type=file_path.suffix.lower(),
            mime_type=mime_type or 'application/octet-stream',
            created_at=str(stat.st_ctime),
            modified_at=str(stat.st_mtime),
            hash_md5=hash_md5.hexdigest()
        )
    
    def _process_pdf(self, file_path: Path) -> str:
        """Process PDF files with fallback to OCR"""
        try:
            # Try pdfplumber first (better for complex layouts)
            return self._process_pdf_with_pdfplumber(file_path)
        except Exception as e1:
            self.logger.warning(f"pdfplumber failed for {file_path}: {e1}")
            
            try:
                # Fallback to PyPDF2
                return self._process_pdf_with_pypdf2(file_path)
            except Exception as e2:
                self.logger.warning(f"PyPDF2 failed for {file_path}: {e2}")
                
                # Final fallback: OCR
                self.logger.info(f"Using OCR for {file_path}")
                return self._process_pdf_with_ocr(file_path)
    
    def _process_pdf_with_pdfplumber(self, file_path: Path) -> str:
        """Process PDF using pdfplumber"""
        import pdfplumber
        
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                except Exception as e:
                    self.logger.warning(f"Error extracting page {page_num}: {e}")
                    continue
        
        return '\\n\\n'.join(text_content)
    
    def _process_pdf_with_pypdf2(self, file_path: Path) -> str:
        """Process PDF using PyPDF2"""
        import PyPDF2
        
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                except Exception as e:
                    self.logger.warning(f"Error extracting page {page_num}: {e}")
                    continue
        
        return '\\n\\n'.join(text_content)
    
    def _process_pdf_with_ocr(self, file_path: Path) -> str:
        """Process PDF using OCR (for scanned documents)"""
        try:
            import pdf2image
            import pytesseract
            from PIL import Image
        except ImportError:
            raise ImportError("pdf2image, pytesseract, and PIL required for OCR")
        
        text_content = []
        
        # Convert PDF pages to images
        try:
            pages = pdf2image.convert_from_path(file_path)
            
            for page_num, page in enumerate(pages):
                try:
                    # Use OCR to extract text
                    text = pytesseract.image_to_string(page, lang='eng')
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    self.logger.warning(f"OCR failed for page {page_num}: {e}")
                    continue
                    
        except Exception as e:
            self.logger.error(f"PDF to image conversion failed: {e}")
            raise
        
        return '\\n\\n'.join(text_content)
    
    def _process_text(self, file_path: Path) -> str:
        """Process plain text files"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode text file {file_path} with any encoding")
    
    def _process_docx(self, file_path: Path) -> str:
        """Process Word documents"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required for processing Word documents")
        
        doc = DocxDocument(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\\n\\n'.join(text_content)
    
    def _process_html(self, file_path: Path) -> str:
        """Process HTML files"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 required for processing HTML")
        
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text content
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def _process_markdown(self, file_path: Path) -> str:
        """Process Markdown files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Basic markdown processing (remove common markdown syntax)
        content = re.sub(r'#{1,6}\\s+', '', content)  # Remove headers
        content = re.sub(r'\\*\\*(.*?)\\*\\*', r'\\1', content)  # Remove bold
        content = re.sub(r'\\*(.*?)\\*', r'\\1', content)  # Remove italic
        content = re.sub(r'\\[([^\\]]+)\\]\\([^\\)]+\\)', r'\\1', content)  # Remove links
        content = re.sub(r'`([^`]+)`', r'\\1', content)  # Remove inline code
        
        return content
    
    def batch_process(self, directory: Union[str, Path], 
                     file_pattern: str = "*",
                     chunking_strategy: str = 'recursive',
                     **chunking_kwargs) -> List[ProcessedDocument]:
        """
        Process multiple documents in a directory
        
        Args:
            directory: Directory containing documents
            file_pattern: Pattern to match files (e.g., "*.pdf")
            chunking_strategy: Strategy for chunking documents
            **chunking_kwargs: Additional arguments for chunking
            
        Returns:
            List of ProcessedDocument objects
        """
        directory = Path(directory)
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        processed_documents = []
        files = list(directory.glob(file_pattern))
        
        self.logger.info(f"Found {len(files)} files to process in {directory}")
        
        for file_path in files:
            if file_path.suffix.lower() in self.supported_formats:
                try:
                    processed_doc = self.process_document(
                        file_path, 
                        chunking_strategy,
                        **chunking_kwargs
                    )
                    processed_documents.append(processed_doc)
                except Exception as e:
                    self.logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        self.logger.info(f"Successfully processed {len(processed_documents)} documents")
        
        return processed_documents

# Global document processor instance
document_processor = DocumentProcessor()
'''

# Need to add missing import for re module
import_fix = '''import re
'''

full_content = import_fix + document_processor_content

with open('multi_doc_rag/core/document_processor.py', 'w') as f:
    f.write(full_content)

print("✅ Document processor created")